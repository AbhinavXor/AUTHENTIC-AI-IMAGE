import asyncio
import json
import logging
import re
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import PurePosixPath
from typing import Any
from zipfile import (
    BadZipFile,
    ZipFile,
    is_zipfile,
)

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from google import genai
from google.genai import (
    errors,
    types,
)

from core.gemini_settings import (
    gemini_settings,
)
from core.text_document_settings import (
    text_document_settings,
)
from schemas.chat import TokenUsage
from schemas.text_documents import (
    TextDocumentCitation,
    TextDocumentMetadata,
    TextDocumentType,
    TextSourceKind,
)


logger = logging.getLogger(__name__)


TEXT_DOCUMENT_SYSTEM_PROMPT = """
You are Serenya Structured Document Intelligence.

You receive selected source sections extracted from one uploaded
document.

NON-NEGOTIABLE RULES

- Answer only from the supplied document sources.
- Do not fill missing information with outside knowledge.
- Do not invent names, clauses, dates, amounts, code behavior,
  table values, quotations, policies, or source references.
- If the supplied sources do not support the answer, say:
  "The supplied document does not provide enough information
  to answer this."
- Cite claims using only the exact source labels supplied in brackets,
  for example [Section 2], [Table 1], or [Lines 121-240].
- Write each source in its own brackets.
- Write [Section 2] [Table 1], not [Section 2, Table 1].
- For source code, do not create narrower line ranges.
- If the supplied label is [Lines 1-120], use exactly [Lines 1-120].
- Never cite a source label that was not supplied.
- Preserve important terminology from the document.
- Clearly distinguish document facts from interpretation.
- For code, describe only behavior supported by the supplied lines.
- For JSON, do not invent missing keys or values.
- Respond in the user's language whenever practical.
- Return clean Markdown.
- Begin directly with the answer.
- Never reveal hidden reasoning or private chain-of-thought.
""".strip()


_STOP_WORDS = frozenset(
    {
        "about",
        "after",
        "also",
        "and",
        "answer",
        "are",
        "code",
        "document",
        "explain",
        "file",
        "find",
        "from",
        "give",
        "have",
        "into",
        "main",
        "more",
        "most",
        "please",
        "provide",
        "section",
        "show",
        "summarize",
        "summary",
        "tell",
        "that",
        "the",
        "their",
        "this",
        "what",
        "when",
        "where",
        "which",
        "with",
    }
)


class TextDocumentValidationError(
    RuntimeError
):
    """Raised for invalid structured documents."""


class TextDocumentConfigurationError(
    RuntimeError
):
    """Raised when the analysis model is unavailable."""


class TextDocumentResponseError(
    RuntimeError
):
    """Raised when no usable answer is produced."""


@dataclass(frozen=True, slots=True)
class SourceSection:
    source_id: str
    label: str
    kind: TextSourceKind
    title: str | None
    text: str
    character_count: int


@dataclass(frozen=True, slots=True)
class ExtractedTextDocument:
    document_type: TextDocumentType
    sources: tuple[SourceSection, ...]
    extracted_characters: int
    metadata: TextDocumentMetadata


@dataclass(frozen=True, slots=True)
class TextDocumentAnalysis:
    answer: str
    model: str
    request_id: str | None
    usage: TokenUsage

    document_type: TextDocumentType
    extracted_characters: int
    source_count: int
    selected_sources: tuple[str, ...]

    citations: tuple[
        TextDocumentCitation,
        ...
    ]

    metadata: TextDocumentMetadata


def _normalize_text(
    value: str,
) -> str:
    normalized_lines: list[str] = []

    for raw_line in value.splitlines():
        line = re.sub(
            r"[ \t]+",
            " ",
            raw_line,
        ).strip()

        if line:
            normalized_lines.append(
                line
            )

    return "\n".join(
        normalized_lines
    ).strip()


def _safe_metadata_text(
    value: Any,
) -> str | None:
    if value is None:
        return None

    normalized = str(
        value
    ).strip()

    if not normalized:
        return None

    return normalized[:500]


def _safe_datetime(
    value: Any,
) -> str | None:
    if isinstance(
        value,
        (
            datetime,
            date,
        ),
    ):
        return value.isoformat()

    return _safe_metadata_text(
        value
    )


def _decode_text_bytes(
    file_bytes: bytes,
) -> str:
    if b"\x00" in file_bytes[:4_096]:
        if file_bytes.startswith(
            (
                b"\xff\xfe",
                b"\xfe\xff",
            )
        ):
            try:
                return file_bytes.decode(
                    "utf-16"
                )
            except UnicodeDecodeError as error:
                raise TextDocumentValidationError(
                    "The text file uses an unsupported encoding."
                ) from error

        raise TextDocumentValidationError(
            "The uploaded file appears to contain binary data."
        )

    try:
        return file_bytes.decode(
            "utf-8-sig"
        )

    except UnicodeDecodeError as error:
        raise TextDocumentValidationError(
            "Text files must use UTF-8 or UTF-16 encoding."
        ) from error


def _split_large_text(
    text: str,
    maximum_characters: int,
) -> list[str]:
    if len(text) <= maximum_characters:
        return [
            text,
        ]

    paragraphs = re.split(
        r"\n{2,}",
        text,
    )

    chunks: list[str] = []
    current: list[str] = []
    current_length = 0

    for paragraph in paragraphs:
        normalized = paragraph.strip()

        if not normalized:
            continue

        if len(normalized) > maximum_characters:
            if current:
                chunks.append(
                    "\n\n".join(current)
                )

                current = []
                current_length = 0

            for position in range(
                0,
                len(normalized),
                maximum_characters,
            ):
                chunks.append(
                    normalized[
                        position:
                        position
                        + maximum_characters
                    ]
                )

            continue

        additional_length = (
            len(normalized)
            + (
                2
                if current
                else 0
            )
        )

        if (
            current
            and current_length
            + additional_length
            > maximum_characters
        ):
            chunks.append(
                "\n\n".join(current)
            )

            current = [
                normalized,
            ]

            current_length = len(
                normalized
            )

        else:
            current.append(
                normalized
            )

            current_length += (
                additional_length
            )

    if current:
        chunks.append(
            "\n\n".join(current)
        )

    return chunks


def _validate_docx_archive(
    file_bytes: bytes,
) -> None:
    stream = BytesIO(
        file_bytes
    )

    if not is_zipfile(stream):
        raise TextDocumentValidationError(
            "The uploaded file is not a valid DOCX archive."
        )

    stream.seek(0)

    try:
        with ZipFile(stream) as archive:
            entries = archive.infolist()

            if (
                len(entries)
                > text_document_settings
                .maximum_docx_archive_entries
            ):
                raise TextDocumentValidationError(
                    "The DOCX archive contains too many entries."
                )

            total_uncompressed_bytes = 0
            entry_names: set[str] = set()

            for entry in entries:
                entry_path = PurePosixPath(
                    entry.filename
                )

                if (
                    entry_path.is_absolute()
                    or ".." in entry_path.parts
                ):
                    raise TextDocumentValidationError(
                        "The DOCX archive contains an unsafe path."
                    )

                if entry.flag_bits & 0x1:
                    raise TextDocumentValidationError(
                        "Encrypted DOCX archives are not supported."
                    )

                total_uncompressed_bytes += (
                    entry.file_size
                )

                if (
                    total_uncompressed_bytes
                    > text_document_settings
                    .maximum_docx_uncompressed_bytes
                ):
                    raise TextDocumentValidationError(
                        "The expanded DOCX exceeds the safety limit."
                    )

                entry_names.add(
                    entry.filename
                )

            required_entries = {
                "[Content_Types].xml",
                "word/document.xml",
            }

            if not required_entries.issubset(
                entry_names
            ):
                raise TextDocumentValidationError(
                    "The archive does not contain a valid Word document."
                )

    except BadZipFile as error:
        raise TextDocumentValidationError(
            "The DOCX archive is corrupted."
        ) from error


def _query_terms(
    prompt: str,
) -> frozenset[str]:
    return frozenset(
        term
        for term in re.findall(
            r"[a-zA-Z0-9][a-zA-Z0-9_-]{2,}",
            prompt.lower(),
        )
        if term not in _STOP_WORDS
    )


def _looks_like_broad_request(
    prompt: str,
) -> bool:
    lowered = prompt.lower()

    broad_terms = (
        "summarize",
        "summary",
        "overview",
        "main points",
        "key points",
        "entire document",
        "whole document",
        "poora document",
        "full document",
    )

    return any(
        term in lowered
        for term in broad_terms
    )


def _source_score(
    source: SourceSection,
    terms: frozenset[str],
) -> float:
    if not terms:
        return 0.0

    searchable = (
        f"{source.label}\n"
        f"{source.title or ''}\n"
        f"{source.text}"
    ).lower()

    score = 0.0

    for term in terms:
        occurrences = searchable.count(
            term
        )

        if occurrences:
            score += min(
                occurrences,
                15,
            ) * 2.0

        if (
            source.title
            and term
            in source.title.lower()
        ):
            score += 6.0

    return score


def _evenly_sample(
    sources: tuple[
        SourceSection,
        ...
    ],
    count: int,
) -> list[SourceSection]:
    if count <= 0 or not sources:
        return []

    if count >= len(sources):
        return list(
            sources
        )

    if count == 1:
        return [
            sources[0],
        ]

    final_index = len(
        sources
    ) - 1

    indexes = {
        round(
            position
            * final_index
            / (count - 1)
        )
        for position in range(
            count
        )
    }

    return [
        sources[index]
        for index in sorted(
            indexes
        )
    ]


class TextDocumentService:
    provider_name = "gemini"

    def __init__(self) -> None:
        self._client: (
            genai.Client
            | None
        ) = None

    def _get_client(
        self,
    ) -> genai.Client:
        if not gemini_settings.api_key:
            raise TextDocumentConfigurationError(
                "Gemini structured-document analysis "
                "is not configured."
            )

        if self._client is None:
            self._client = genai.Client(
                api_key=gemini_settings.api_key,
                http_options=types.HttpOptions(
                    timeout=int(
                        gemini_settings.timeout_seconds
                        * 1000
                    ),
                ),
            )

        return self._client

    @staticmethod
    def _model_candidates(
    ) -> tuple[str, ...]:
        return tuple(
            dict.fromkeys(
                model
                for model in (
                    gemini_settings.quality_model,
                    gemini_settings.fallback_model,
                    gemini_settings.preview_model,
                    gemini_settings.fast_model,
                )
                if model
            )
        )

    @staticmethod
    def _build_docx_sources(
        file_bytes: bytes,
    ) -> ExtractedTextDocument:
        _validate_docx_archive(
            file_bytes
        )

        try:
            document = Document(
                BytesIO(
                    file_bytes
                )
            )

        except Exception as error:
            raise TextDocumentValidationError(
                "The DOCX document could not be opened."
            ) from error

        raw_blocks: list[
            tuple[
                TextSourceKind,
                str | None,
                str,
            ]
        ] = []

        current_heading: (
            str
            | None
        ) = None

        current_parts: list[str] = []

        def flush_section() -> None:
            nonlocal current_parts

            if not current_parts:
                return

            text = "\n\n".join(
                current_parts
            ).strip()

            if text:
                raw_blocks.append(
                    (
                        "section",
                        current_heading,
                        text,
                    )
                )

            current_parts = []

        for item in (
            document.iter_inner_content()
        ):
            if isinstance(
                item,
                Paragraph,
            ):
                paragraph_text = (
                    _normalize_text(
                        item.text
                    )
                )

                if not paragraph_text:
                    continue

                style_name = (
                    item.style.name
                    if item.style
                    else ""
                )

                if style_name.lower().startswith(
                    "heading"
                ):
                    flush_section()

                    current_heading = (
                        paragraph_text
                    )

                    current_parts = [
                        paragraph_text,
                    ]

                else:
                    current_parts.append(
                        paragraph_text
                    )

            elif isinstance(
                item,
                Table,
            ):
                flush_section()

                table_rows: list[str] = []

                for row in item.rows:
                    cells = [
                        _normalize_text(
                            cell.text
                        )
                        for cell in row.cells
                    ]

                    if any(cells):
                        table_rows.append(
                            " | ".join(
                                cells
                            )
                        )

                table_text = "\n".join(
                    table_rows
                ).strip()

                if table_text:
                    raw_blocks.append(
                        (
                            "table",
                            current_heading,
                            table_text,
                        )
                    )

                current_heading = None

        flush_section()

        sources: list[
            SourceSection
        ] = []

        section_number = 0
        table_number = 0

        for (
            kind,
            title,
            raw_text,
        ) in raw_blocks:
            chunks = _split_large_text(
                raw_text,
                text_document_settings
                .maximum_section_characters,
            )

            for chunk in chunks:
                if kind == "table":
                    table_number += 1

                    label = (
                        f"Table {table_number}"
                    )

                    source_id = (
                        f"table-{table_number}"
                    )

                else:
                    section_number += 1

                    label = (
                        f"Section {section_number}"
                    )

                    source_id = (
                        f"section-{section_number}"
                    )

                sources.append(
                    SourceSection(
                        source_id=source_id,
                        label=label,
                        kind=kind,
                        title=title,
                        text=chunk,
                        character_count=len(
                            chunk
                        ),
                    )
                )

        if (
            len(sources)
            > text_document_settings
            .maximum_sections
        ):
            raise TextDocumentValidationError(
                "The DOCX contains too many document sections."
            )

        extracted_characters = sum(
            source.character_count
            for source in sources
        )

        if (
            extracted_characters
            < text_document_settings
            .minimum_usable_characters
        ):
            raise TextDocumentValidationError(
                "The DOCX does not contain enough readable text."
            )

        core = document.core_properties

        metadata = TextDocumentMetadata(
            title=_safe_metadata_text(
                core.title
            ),
            author=_safe_metadata_text(
                core.author
            ),
            subject=_safe_metadata_text(
                core.subject
            ),
            keywords=_safe_metadata_text(
                core.keywords
            ),
            created=_safe_datetime(
                core.created
            ),
            modified=_safe_datetime(
                core.modified
            ),
        )

        return ExtractedTextDocument(
            document_type="docx",
            sources=tuple(
                sources
            ),
            extracted_characters=(
                extracted_characters
            ),
            metadata=metadata,
        )

    @staticmethod
    def _build_code_sources(
        text: str,
    ) -> tuple[
        SourceSection,
        ...
    ]:
        lines = text.splitlines()

        chunk_size = (
            text_document_settings
            .code_chunk_lines
        )

        overlap = (
            text_document_settings
            .code_chunk_overlap_lines
        )

        step = max(
            1,
            chunk_size - overlap,
        )

        sources: list[
            SourceSection
        ] = []

        for start_index in range(
            0,
            len(lines),
            step,
        ):
            end_index = min(
                start_index + chunk_size,
                len(lines),
            )

            numbered_lines = [
                f"{line_number}: {lines[line_number - 1]}"
                for line_number in range(
                    start_index + 1,
                    end_index + 1,
                )
            ]

            chunk = "\n".join(
                numbered_lines
            ).strip()

            if not chunk:
                continue

            first_line = (
                start_index + 1
            )

            final_line = end_index

            label = (
                f"Lines {first_line}-{final_line}"
            )

            sources.append(
                SourceSection(
                    source_id=(
                        f"lines-{first_line}-{final_line}"
                    ),
                    label=label,
                    kind="lines",
                    title=None,
                    text=chunk,
                    character_count=len(
                        chunk
                    ),
                )
            )

            if end_index >= len(lines):
                break

        return tuple(
            sources
        )

    @staticmethod
    def _build_text_sources(
        *,
        file_bytes: bytes,
        extension: str,
    ) -> ExtractedTextDocument:
        decoded = _decode_text_bytes(
            file_bytes
        )

        if extension in (
            text_document_settings
            .json_extensions
        ):
            try:
                parsed_json = json.loads(
                    decoded
                )

            except json.JSONDecodeError as error:
                raise TextDocumentValidationError(
                    "The JSON document is malformed."
                ) from error

            decoded = json.dumps(
                parsed_json,
                ensure_ascii=False,
                indent=2,
            )

            document_type: TextDocumentType = (
                "json"
            )

        elif extension in (
            text_document_settings
            .markdown_extensions
        ):
            document_type = "markdown"

        elif extension in (
            text_document_settings
            .source_code_extensions
        ):
            document_type = "source_code"

        else:
            document_type = "text"

        normalized = decoded.strip()

        if (
            len(normalized)
            < text_document_settings
            .minimum_usable_characters
        ):
            raise TextDocumentValidationError(
                "The document does not contain enough readable text."
            )

        if document_type == "source_code":
            sources = (
                TextDocumentService
                ._build_code_sources(
                    normalized
                )
            )

        else:
            chunks = _split_large_text(
                normalized,
                text_document_settings
                .maximum_section_characters,
            )

            sources = tuple(
                SourceSection(
                    source_id=(
                        f"section-{index}"
                    ),
                    label=(
                        f"Section {index}"
                    ),
                    kind="section",
                    title=None,
                    text=chunk,
                    character_count=len(
                        chunk
                    ),
                )
                for index, chunk in enumerate(
                    chunks,
                    start=1,
                )
            )

        if not sources:
            raise TextDocumentValidationError(
                "No readable document sections were found."
            )

        if (
            len(sources)
            > text_document_settings
            .maximum_sections
        ):
            raise TextDocumentValidationError(
                "The document contains too many sections."
            )

        return ExtractedTextDocument(
            document_type=document_type,
            sources=sources,
            extracted_characters=sum(
                source.character_count
                for source in sources
            ),
            metadata=TextDocumentMetadata(),
        )

    @staticmethod
    def _extract_sync(
        *,
        file_bytes: bytes,
        extension: str,
    ) -> ExtractedTextDocument:
        if extension in (
            text_document_settings
            .docx_extensions
        ):
            return (
                TextDocumentService
                ._build_docx_sources(
                    file_bytes
                )
            )

        return (
            TextDocumentService
            ._build_text_sources(
                file_bytes=file_bytes,
                extension=extension,
            )
        )

    @staticmethod
    def _select_sources(
        document: ExtractedTextDocument,
        prompt: str,
    ) -> tuple[
        SourceSection,
        ...
    ]:
        sources = document.sources

        maximum_sources = (
            text_document_settings
            .maximum_selected_sections
        )

        if len(sources) <= maximum_sources:
            return sources

        selected: dict[
            str,
            SourceSection,
        ] = {}

        def add_source(
            source: SourceSection,
        ) -> None:
            if (
                len(selected)
                < maximum_sources
            ):
                selected.setdefault(
                    source.source_id,
                    source,
                )

        for source in sources[:2]:
            add_source(
                source
            )

        if _looks_like_broad_request(
            prompt
        ):
            for source in _evenly_sample(
                sources,
                maximum_sources,
            ):
                add_source(
                    source
                )

        else:
            terms = _query_terms(
                prompt
            )

            ranked = sorted(
                sources,
                key=lambda source: (
                    _source_score(
                        source,
                        terms,
                    ),
                    source.character_count,
                ),
                reverse=True,
            )

            for source in ranked:
                if (
                    _source_score(
                        source,
                        terms,
                    )
                    <= 0
                ):
                    continue

                add_source(
                    source
                )

            remaining = (
                maximum_sources
                - len(selected)
            )

            for source in _evenly_sample(
                sources,
                remaining,
            ):
                add_source(
                    source
                )

        add_source(
            sources[-1]
        )

        return tuple(
            source
            for source in sources
            if source.source_id
            in selected
        )

    @staticmethod
    def _build_context(
        sources: tuple[
            SourceSection,
            ...
        ],
    ) -> str:
        sections: list[str] = []
        current_characters = 0

        for source in sources:
            title_suffix = (
                f" | {source.title}"
                if source.title
                else ""
            )

            section = (
                f"\n===== SOURCE: "
                f"{source.label}"
                f"{title_suffix} =====\n"
                f"{source.text}\n"
            )

            if (
                current_characters
                + len(section)
                > text_document_settings
                .maximum_context_characters
            ):
                remaining = (
                    text_document_settings
                    .maximum_context_characters
                    - current_characters
                )

                if remaining > 300:
                    sections.append(
                        section[:remaining]
                    )

                break

            sections.append(
                section
            )

            current_characters += len(
                section
            )

        return "".join(
            sections
        ).strip()

    @staticmethod
    def _response_text(
        response: Any,
    ) -> str:
        try:
            return (
                getattr(
                    response,
                    "text",
                    "",
                )
                or ""
            ).strip()

        except Exception:
            return ""

    @staticmethod
    def _usage(
        response: Any,
    ) -> TokenUsage:
        metadata = getattr(
            response,
            "usage_metadata",
            None,
        )

        if metadata is None:
            return TokenUsage()

        prompt_tokens = int(
            getattr(
                metadata,
                "prompt_token_count",
                0,
            )
            or 0
        )

        completion_tokens = int(
            getattr(
                metadata,
                "candidates_token_count",
                0,
            )
            or getattr(
                metadata,
                "output_token_count",
                0,
            )
            or 0
        )

        total_tokens = int(
            getattr(
                metadata,
                "total_token_count",
                0,
            )
            or (
                prompt_tokens
                + completion_tokens
            )
        )

        return TokenUsage(
            prompt_tokens=prompt_tokens,
            completion_tokens=(
                completion_tokens
            ),
            total_tokens=total_tokens,
        )

    @staticmethod
    def _status_code(
        error: Exception,
    ) -> int | None:
        raw_code = getattr(
            error,
            "code",
            None,
        )

        try:
            return int(
                raw_code
            )

        except (
            TypeError,
            ValueError,
        ):
            return None

    @staticmethod
    def _canonicalize_answer_citations(
        *,
        answer: str,
        sources: tuple[
            SourceSection,
            ...
        ],
    ) -> str:
        """
        Convert model-produced citations into exact,
        supplied source labels.

        Examples:

        [Section 2, Table 1]
        becomes
        [Section 2] [Table 1]

        [Lines 7-8]
        becomes the supplied enclosing source,
        such as
        [Lines 1-13]
        """

        line_sources: list[
            tuple[
                SourceSection,
                int,
                int,
            ]
        ] = []

        for source in sources:
            line_match = re.fullmatch(
                r"Lines\s+(\d+)\s*[-–]\s*(\d+)",
                source.label,
                flags=re.IGNORECASE,
            )

            if line_match is None:
                continue

            source_start = int(
                line_match.group(1)
            )

            source_end = int(
                line_match.group(2)
            )

            if source_start > source_end:
                source_start, source_end = (
                    source_end,
                    source_start,
                )

            line_sources.append(
                (
                    source,
                    source_start,
                    source_end,
                )
            )

        def replace_bracket_group(
            match: re.Match[str],
        ) -> str:
            raw_content = (
                match.group(1).strip()
            )

            lowered_content = (
                raw_content.lower()
            )

            canonical_labels: list[
                str
            ] = []

            # Handle exact and combined citations such as:
            # [Section 2, Table 1]
            for source in sources:
                source_label = (
                    source.label
                )

                escaped_label = re.escape(
                    source_label.lower()
                )

                if re.search(
                    (
                        r"(?<![a-zA-Z0-9])"
                        + escaped_label
                        + r"(?![a-zA-Z0-9])"
                    ),
                    lowered_content,
                ):
                    if (
                        source_label
                        not in canonical_labels
                    ):
                        canonical_labels.append(
                            source_label
                        )

            # Handle unsupported narrower model ranges,
            # such as [Lines 7-8], by mapping them to
            # the enclosing supplied source range.
            if not canonical_labels:
                requested_lines = re.fullmatch(
                    r"Lines\s+(\d+)\s*[-–]\s*(\d+)",
                    raw_content,
                    flags=re.IGNORECASE,
                )

                if requested_lines is not None:
                    requested_start = int(
                        requested_lines.group(1)
                    )

                    requested_end = int(
                        requested_lines.group(2)
                    )

                    if (
                        requested_start
                        > requested_end
                    ):
                        (
                            requested_start,
                            requested_end,
                        ) = (
                            requested_end,
                            requested_start,
                        )

                    for (
                        source,
                        source_start,
                        source_end,
                    ) in line_sources:
                        if (
                            source_start
                            <= requested_start
                            and requested_end
                            <= source_end
                        ):
                            canonical_labels.append(
                                source.label
                            )

                            break

            if not canonical_labels:
                return match.group(0)

            return " ".join(
                f"[{label}]"
                for label in canonical_labels
            )

        return re.sub(
            r"\[([^\[\]\n]{1,180})\]",
            replace_bracket_group,
            answer,
        )

    @staticmethod
    def _extract_citations(
        *,
        answer: str,
        sources: tuple[
            SourceSection,
            ...
        ],
    ) -> tuple[
        TextDocumentCitation,
        ...
    ]:
        citations: list[
            TextDocumentCitation
        ] = []

        lowered_answer = (
            answer.lower()
        )

        for source in sources:
            citation_token = (
                f"[{source.label}]"
            )

            if (
                citation_token.lower()
                not in lowered_answer
            ):
                continue

            citations.append(
                TextDocumentCitation(
                    source_id=(
                        source.source_id
                    ),
                    label=source.label,
                    kind=source.kind,
                )
            )

        return tuple(
            citations
        )

    async def analyze(
        self,
        *,
        file_bytes: bytes,
        extension: str,
        prompt: str,
    ) -> TextDocumentAnalysis:
        document = await asyncio.to_thread(
            self._extract_sync,
            file_bytes=file_bytes,
            extension=extension,
        )

        selected_sources = (
            self._select_sources(
                document,
                prompt,
            )
        )

        source_labels = tuple(
            source.label
            for source in selected_sources
        )

        context = self._build_context(
            selected_sources
        )

        request_content = f"""
USER REQUEST

{prompt}

DOCUMENT TYPE

{document.document_type}

AVAILABLE SOURCE LABELS

{", ".join(source_labels)}

SUPPLIED DOCUMENT CONTENT

{context}

Answer using only the supplied content.
Use exact inline source citations such as [Section 1],
[Table 1], or [Lines 1-120].
""".strip()

        last_error: (
            Exception
            | None
        ) = None

        for model in (
            self._model_candidates()
        ):
            try:
                response = (
                    await self
                    ._get_client()
                    .aio
                    .models
                    .generate_content(
                        model=model,
                        contents=request_content,
                        config=types.GenerateContentConfig(
                            system_instruction=(
                                TEXT_DOCUMENT_SYSTEM_PROMPT
                            ),
                            temperature=0.1,
                            max_output_tokens=3_000,
                        ),
                    )
                )

                answer = (
                    self._response_text(
                        response
                    )
                )

                if not answer:
                    raise TextDocumentResponseError(
                        "The structured-document model "
                        "returned an empty answer."
                    )

                answer = (
                    self._canonicalize_answer_citations(
                        answer=answer,
                        sources=selected_sources,
                    )
                )

                citations = (
                    self._extract_citations(
                        answer=answer,
                        sources=selected_sources,
                    )
                )

                if not citations:
                    reviewed_sources = (
                        ", ".join(
                            f"[{source.label}]"
                            for source
                            in selected_sources[:6]
                        )
                    )

                    answer = (
                        answer.rstrip()
                        + "\n\n"
                        + "**Sources reviewed:** "
                        + reviewed_sources
                    )

                    citations = tuple(
                        TextDocumentCitation(
                            source_id=(
                                source.source_id
                            ),
                            label=source.label,
                            kind=source.kind,
                        )
                        for source
                        in selected_sources[:6]
                    )

                return TextDocumentAnalysis(
                    answer=answer,
                    model=model,
                    request_id=getattr(
                        response,
                        "response_id",
                        None,
                    ),
                    usage=self._usage(
                        response
                    ),
                    document_type=(
                        document.document_type
                    ),
                    extracted_characters=(
                        document
                        .extracted_characters
                    ),
                    source_count=len(
                        document.sources
                    ),
                    selected_sources=(
                        source_labels
                    ),
                    citations=citations,
                    metadata=document.metadata,
                )

            except Exception as error:
                last_error = error

                logger.warning(
                    "Structured document model failed: "
                    "model=%s type=%s status=%s",
                    model,
                    type(error).__name__,
                    self._status_code(
                        error
                    ),
                )

                if isinstance(
                    error,
                    errors.APIError,
                ):
                    status_code = (
                        self._status_code(
                            error
                        )
                    )

                    if status_code in {
                        401,
                        403,
                    }:
                        raise TextDocumentConfigurationError(
                            "Gemini document credentials "
                            "are invalid or unauthorized."
                        ) from error

                    if status_code not in {
                        404,
                        408,
                        429,
                        500,
                        502,
                        503,
                        504,
                    }:
                        raise TextDocumentResponseError(
                            "Gemini rejected the "
                            "structured-document request."
                        ) from error

                elif not isinstance(
                    error,
                    TextDocumentResponseError,
                ):
                    raise TextDocumentResponseError(
                        "The structured-document request "
                        "could not be completed."
                    ) from error

        raise TextDocumentResponseError(
            "No configured model produced a usable "
            "structured-document answer."
        ) from last_error
